
import argparse
import os
parser = argparse.ArgumentParser(description='Colorize GreyScale Image')
parser.add_argument('--input', help='Please give path to image')
args = parser.parse_args()

if args.input==None:
    print('Please give the input greyscale image name.')
    print('Usage example: python fill_color.py --input greyscaleImage.png')
    exit()

if os.path.isfile(args.input)==0:
    print('Input file does not exist')
    exit()

filename = args.input

from get_data import get_data

raw_text = get_data.getText(filename)

import spacy
#import en_core_web_sm
nlp = spacy.load("en_core_web_sm")

paras = [i.replace('\t',' ') for i in raw_text.split('\n') if i!='']

inp_to_spacy = " ".join(paras) # create string from paras list

doc = nlp(inp_to_spacy) # a spacy doc object it has everything

# getting sentences out of doc
sentences = [sentence for idno, sentence in enumerate(doc.sents)]

import re
sentences = [re.sub("…", "", str(sentence)) for sentence in sentences]

sentences = [re.sub("[.][.]+", "", str(sentence)) for sentence in sentences]

# again we input the data in spacy

input_to_spacy = " ".join(sentences)
doc = nlp(input_to_spacy)


tokens = [token for token in doc if not nlp.vocab[str(token)].is_stop]

# still it contains tokens that we dont need like , . - etc
# removing them manually
from stop import stop_words_class
stop_words = stop_words_class.stop_words
tokens = [token for token in doc if str(token) not in stop_words]


# ## Result from the above cells:
# ## 1. `sentences`
# ## 2. `tokens`


import numpy as np
import pandas as pd
import nltk
nltk.download('punkt') # one time execution
import re


# ## Need to download dataset
# http://downloads.cs.stanford.edu/nlp/data/glove.6B.zip

word_embeddings = {}

f = open('../glove.6B/glove.6B.100d.txt', encoding='utf-8')
for line in f:
    values = line.split()
    word = values[0]
    coefs = np.asarray(values[1:], dtype='float32')
    word_embeddings[word] = coefs
f.close()


#    create sentence vectors
sentence_vectors = []
for i in sentences:
    if len(i) != 0:
        v = sum([word_embeddings.get(w, np.zeros((100,))) for w in i.split()])/(len(i.split())+0.001)
    else:
        v = np.zeros((100,))
    sentence_vectors.append(v)



sentence_vectors

# similarity matrix
sim_mat = np.zeros([len(sentences), len(sentences)])


from sklearn.metrics.pairwise import cosine_similarity


for i in range(len(sentences)):
    for j in range(len(sentences)):
        if i != j:
            sim_mat[i][j] = cosine_similarity(sentence_vectors[i].reshape(1,100), sentence_vectors[j].reshape(1,100))[0,0]

import networkx as nx

nx_graph = nx.from_numpy_matrix(sim_mat)
scores = nx.pagerank(nx_graph)


ranked_sentences = sorted(((scores[i],s) for i,s in enumerate(sentences)), reverse=True)


summary = []
for i in range(12):
    summary.append(ranked_sentences[i][1])


ordered_summary = []
for i in sentences:
    if i in summary:
        ordered_summary.append(i)
import docx
d = docx.Document()
d.add_heading('Summary', 0)

for i in ordered_summary:
    d.add_paragraph(
        i, style='List Bullet'
    )

d.save(filename[:len(filename)-5]+'_summarized.docx')

