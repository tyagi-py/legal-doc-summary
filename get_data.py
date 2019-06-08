class get_data:
    
    def getText(filename):
        import docx
        from tika import parser
        if filename[len(filename)-3:] == "pdf":
            raw = parser.from_file('raw_data/selab.pdf')
            return raw['content']
        else:
            doc = docx.Document(filename)
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
        return '\n'.join(fullText)

    def getTable(filename, n =0):
        from docx.api import Document
        document = Document(filename)
        table = document.tables[n]
        data = []
        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            if i == 0:
                keys = tuple(text)
                continue
            row_data = dict(zip(keys, text))
            data.append(row_data)
        print(data)